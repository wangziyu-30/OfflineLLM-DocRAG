from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional
import json
import io
# 处理Word文档依赖
import docx
import win32com.client

# 导入你的项目核心模块
from rag.vector_store import VectorStoreManager
from rag.config import get_semantic_text_splitter
from rag.chain import rag_chain_with_memory
from langchain_core.documents import Document

# ========== 初始化 ==========
app = FastAPI(title="RAG知识库问答系统", version="1.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 加载本地向量库 + 分块器
vector_store = VectorStoreManager()
vector_store.load_vector_store()
text_splitter = get_semantic_text_splitter()


# ========== 基础模型 ==========
class BaseResponse(BaseModel):
    code: int = 200
    message: str = "success"
    data: Optional[dict] = None


class ChatRequest(BaseModel):
    session_id: str = "user1"
    question: str


# ========== 首页 ==========
@app.get("/", response_model=BaseResponse)
async def index():
    return BaseResponse(message="RAG服务已启动 | 支持 TXT/MD/DOC/DOCX 上传")


# ========== 读取不同格式文件的工具函数 ==========
def read_docx(file_bytes: bytes) -> str:
    """读取DOCX文件"""
    doc = docx.Document(io.BytesIO(file_bytes))
    return "\n".join([para.text for para in doc.paragraphs])


def read_doc(file_path: str) -> str:
    """Windows下读取老版DOC文件"""
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(file_path)
        content = doc.Content.Text
        doc.Close()
        word.Quit()
        return content
    except Exception as e:
        raise Exception(f"DOC文件读取失败：{str(e)}")


# ========== 🔥 全格式上传接口（TXT/MD/DOC/DOCX） ==========
@app.post("/upload", response_model=BaseResponse)
async def upload_file(file: UploadFile = File(...)):
    # 支持所有格式
    ext = file.filename.split(".")[-1].lower()
    support_exts = ["txt", "md", "markdown", "doc", "docx"]
    if ext not in support_exts:
        raise HTTPException(status_code=400, detail="仅支持：TXT/MD/DOC/DOCX 文件")

    try:
        file_bytes = await file.read()
        content = ""

        # 1. 读取文本类文件
        if ext in ["txt", "md", "markdown"]:
            try:
                content = file_bytes.decode("utf-8")
            except:
                content = file_bytes.decode("gbk")

        # 2. 读取DOCX文件
        elif ext == "docx":
            content = read_docx(file_bytes)

        # 3. 读取DOC文件（老版Word）
        elif ext == "doc":
            import tempfile
            with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as temp_file:
                temp_file.write(file_bytes)
                temp_path = temp_file.name
            content = read_doc(temp_path)

        # 文档分块
        doc = Document(page_content=content, metadata={"filename": file.filename})
        split_docs = text_splitter.split_documents([doc])

        # 添加到知识库
        vector_store.add_documents(split_docs)

        return BaseResponse(
            data={
                "文件名": file.filename,
                "分块数": len(split_docs),
                "当前知识库总条数": vector_store.get_document_count()
            }
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"上传失败：{str(e)}")


# ========== 清空知识库 ==========
@app.delete("/clear", response_model=BaseResponse)
async def clear_knowledge_base():
    try:
        vector_store.docs = []
        vector_store.vector_store = None
        return BaseResponse(message="知识库已清空成功！")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"清空失败：{str(e)}")


# ========== 问答接口 ==========
@app.post("/chat", response_model=BaseResponse)
async def chat(request: ChatRequest):
    if not request.question:
        raise HTTPException(status_code=400, detail="请输入问题")

    answer = rag_chain_with_memory.invoke({"question": request.question})
    return BaseResponse(data={"answer": answer})


# ========== 启动服务 ==========
if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="127.0.0.1", port=8000)